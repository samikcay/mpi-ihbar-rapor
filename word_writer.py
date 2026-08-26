# -*- coding: utf-8 -*-
"""
Word raporundaki (Sanal Kumar ve Yasa Disi Bahisle Mucadele Calismalari)
SAP'tan turetilebilen sayilari gunceller.

NEDEN REGEX / RUN BIRLESTIRME?
Word bir cumleyi duzeltme gecmisi ve yazim denetimi yuzunden parcalara
(run) boler. Ornegin "123.456" belgede su sekilde durur:
    run1='1'  run2='23'  run3='.'  run4='456'
Bu yuzden "123.456'yi bul degistir" diyen basit bir arama HICBIR SEY bulamaz.
Cozum: paragrafin tam metnini alip regex ile degistirmek, sonra metni ilk
run'a yazip digerlerini bosaltmak. Bu paragraflarin tum run'lari ayni
bicimde oldugu icin (Times New Roman 11.5) bicim kaybi olmaz - bu
kontrol edilerek dogrulanmistir.

SADECE SAP'TAN KESIN TURETILEBILEN sayilar degistirilir. Turetilemeyenler
(erisime engellenen, EEK, mahkeme karari vb.) ELLENMEZ, uyari olarak
raporlanir.
"""

import os
import re
import shutil
import datetime

import tarih

try:
    import docx
    DOCX_VAR = True
except ImportError:
    DOCX_VAR = False


def _kisa(e):
    """Hata metnini kisaltir (hatalar modulunu kullanir)."""
    try:
        import hatalar
        return hatalar._kisa_hata(e)
    except Exception:
        return str(e)


class WordHatasi(Exception):
    """Word guncellemesi sirasinda beklenmeyen durum."""


# --- Yardimcilar ------------------------------------------------------------

def bicimle(sayi):
    """12345 -> '12.345' (Turkce binlik ayraci)"""
    return "{:,}".format(int(sayi)).replace(",", ".")


def _paragraf_metni_degistir(p, yeni_metin):
    """
    Paragrafin metnini korunakli sekilde degistirir.

    Metni ilk run'a yazar, kalan run'lari bosaltir (silmez - silmek
    bookmark/alan gibi seyleri bozabilir). Bu paragraflarda tum run'lar
    ayni bicimde oldugu icin gorunum degismez.
    """
    if not p.runs:
        return False
    p.runs[0].text = yeni_metin
    for r in p.runs[1:]:
        r.text = ""
    return True


def _sayi_deseni(sayi):
    """
    Belgedeki bir sayiyi (nokta ayracli veya ayracsiz) yakalayan desen.
    '66.411' ve '66411' ikisini de bulur.
    """
    s = str(int(sayi))
    parcali = bicimle(sayi)
    return re.compile(r"(?<![\d.])(?:%s|%s)(?![\d])"
                      % (re.escape(parcali), re.escape(s)))


# --- Guncelleme kurallari ---------------------------------------------------

def _gun_ay_metni(d):
    """date -> '17 Ağustos' (metin icindeki '01 Ocak-14 Ağustos' bicimi)"""
    return "%02d %s" % (d.day, tarih.AYLAR[d.month])


def hazirla_degerler(veri, bas, bit):
    """
    SAP verisinden Word'e yazilacak degerleri hesaplar.

    veri: rapor.calistir icindeki {'icerik':..., 'durum':..., 'ulke':...}
    Doner: sozluk
    """
    icerik = dict(veri["icerik"]["items"])
    ulkeler = veri["ulke"]["items"]

    toplam = sum(v for _, v in ulkeler)

    # Etiketler SAP'ta aksansiz gelebilir; esnek arama
    def bul(sozluk, *anahtarlar):
        for k, v in sozluk.items():
            sade = (k.lower().replace("ı", "i").replace("İ", "i")
                     .replace("ş", "s").replace("ğ", "g").replace("ü", "u")
                     .replace("ö", "o").replace("ç", "c"))
            for a in anahtarlar:
                if a in sade:
                    return v
        return None

    direk = bul(icerik, "direk")
    reklam = bul(icerik, "tanitim", "reklam")

    return {
        "toplam": toplam,
        "direk": direk,
        "reklam": reklam,
        "ulke_sayisi": len(ulkeler),
        "ulkeler": ulkeler,
        "bas": bas,
        "bit": bit,
    }


def _ulke_yuzdeleri(ulkeler, adet=10):
    """
    Ilk N ulkenin yuvarlanmis yuzdesi + kalan dilim.

    ONEMLI: Her yuzde ayri ayri yuvarlanirsa toplam 100 olmayabilir
    (orn. 34+30+5+5+5+3+3+3+3+1+9 = 101). Belgedeki kullanim toplamin
    100 olmasidir; bu yuzden yuvarlama farkini KALAN dilim emer -
    Excel'deki 'Diğer' hucresiyle ayni mantik.
    """
    T = sum(v for _, v in ulkeler)
    if not T:
        return [], 0
    ilk = ulkeler[:adet]
    yuzdeler = [(ad, int(round(v * 100.0 / T))) for ad, v in ilk]
    # Kalan = 100 - ilk N'in yuvarlanmis toplami (fark burada erir)
    kalan = 100 - sum(y for _, y in yuzdeler)
    if kalan < 0:
        kalan = 0
    return yuzdeler, kalan



# ---------------------------------------------------------------------------
# Excel'den okunan degerler
# ---------------------------------------------------------------------------
# Bu sayilar SAP ciktilarinda YOKTUR; calisma kitabinda personelin elle
# doldurdugu kaynak sekmelerde durur. Word'e oradan tasinir.
#
#   anahtar -> (sekme, hucre, aciklama)
EXCEL_KAYNAKLARI = {
    "kumulatif_tespit":  ("Yıllar Bazlı", "B25", "2006-... toplam tespit"),
    "kumulatif_engel":   ("Yıllar Bazlı", "E25", "2006-... erisime engellenen"),
    "mobil":             ("Mobil Uygulamalar", "B4", "mobil uygulama"),
    "sosyal_basvuru":    ("Sosyal Medya Hesapları", "B4", "sosyal medya basvuru"),
    "sosyal_engel":      ("Sosyal Medya Hesapları", "C4", "sosyal medya engellenen"),
    "turetilmis_site":   ("Suç Duyuruları", "C26", "turetilmis site"),
    "kok_site":          ("Suç Duyuruları", "D26", "kok site"),
    "odeme_kurulusu":    ("Ödeme Kuruluşları", "B6", "odeme kurulusu"),
    "hat_850":           ("850'li Hatlar", "B4", "850'li hat"),
    "iban":              ("Faaliyet Cetveli", "H9", "IBAN"),
}


def excel_degerleri_oku(xlsx_yolu):
    """
    Calisma kitabindan Word'e tasinacak sayilari okur.

    Formullerin HESAPLANMIS degerini okur (data_only=True). Dosya Excel
    tarafindan kaydedildigi icin bu degerler gunceldir.

    Doner: (degerler, eksikler)
      degerler : {anahtar: sayi}
      eksikler : [(aciklama, sebep), ...]  -> okunamayanlar SESSIZ GECILMEZ,
                 cagirana bildirilir ki kullanici eski deger kaldigini bilsin.
    """
    try:
        import openpyxl
    except ImportError:
        return {}, [("tümü", "openpyxl kurulu değil")]
    try:
        wb = openpyxl.load_workbook(xlsx_yolu, data_only=True)
    except Exception as e:
        return {}, [("tümü", "çalışma kitabı okunamadı: %s" % e)]

    # Sekme adlarini aksan/bosluk duyarsiz esle
    def sade(x):
        return " ".join((x or "").lower().replace("ı", "i")
                        .replace("İ", "i").replace("ş", "s")
                        .replace("ğ", "g").replace("ü", "u")
                        .replace("ö", "o").replace("ç", "c").split())

    sekmeler = {sade(ws.title): ws for ws in wb.worksheets}

    out = {}
    eksik = []
    for anahtar, (sekme, hucre, aciklama) in sorted(EXCEL_KAYNAKLARI.items()):
        ws = sekmeler.get(sade(sekme))
        if ws is None:
            eksik.append((aciklama, "'%s' sekmesi bulunamadı" % sekme))
            continue
        try:
            v = ws[hucre].value
        except Exception:
            v = None
        if isinstance(v, (int, float)) and v > 0:
            out[anahtar] = int(v)
        else:
            eksik.append((aciklama,
                          "%s!%s boş veya sayı değil" % (sekme, hucre)))
    return out, eksik


# --- Ana islem --------------------------------------------------------------

def yedek_al(yol):
    """Word dosyasinin zaman damgali yedegini alir."""
    klasor = os.path.join(os.path.dirname(os.path.abspath(yol)), "yedek")
    if not os.path.isdir(klasor):
        os.makedirs(klasor)
    damga = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    kok, uz = os.path.splitext(os.path.basename(yol))
    hedef = os.path.join(klasor, "%s_%s%s" % (kok, damga, uz))
    shutil.copy2(yol, hedef)
    return hedef


def guncelle(docx_yolu, veri, bas, bit, xlsx_yolu=None, log=None):
    """
    Word raporundaki SAP kaynakli sayilari gunceller.

    Doner: {'degisen': [...], 'elle': [...], 'yedek': yol}
    """
    if not DOCX_VAR:
        raise WordHatasi(
            "python-docx kurulu degil.\n"
            "Kurmak icin:  pip install python-docx")
    if not os.path.isfile(docx_yolu):
        raise WordHatasi("Word dosyasi bulunamadi: %s" % docx_yolu)

    d = hazirla_degerler(veri, bas, bit)
    # Calisma kitabindan gelen (SAP'ta olmayan) sayilar
    if xlsx_yolu:
        d["excel"], excel_eksik = excel_degerleri_oku(xlsx_yolu)
    else:
        d["excel"], excel_eksik = {}, []
    yedek = yedek_al(docx_yolu)

    try:
        belge = docx.Document(docx_yolu)
    except Exception as e:
        raise WordHatasi(
            "Word dosyasi acilamadi: %s\n"
            "Dosya bozuk olabilir veya .docx bicimi degildir.\n"
            "(Eski .doc dosyalari desteklenmez; Word'de .docx olarak kaydedin.)"
            % _kisa(e))
    degisen = []
    elle = []

    donem_nokta = "%s-%s" % (tarih.nokta(bas), tarih.nokta(bit))

    for sira, p in enumerate(belge.paragraphs):
        metin = p.text
        if not metin.strip():
            continue
        yeni = metin

        # 1) Baslik ve metin icindeki '01 Ocak-14 Ağustos' donemleri
        #    Ayirac (- veya – , bosluklu/bosluksuz) OLDUGU GIBI korunur;
        #    sadece gun ve ay adlari degistirilir.
        yeni = re.sub(
            r"(\d{1,2}\s+(?:Ocak|Şubat|Mart|Nisan|Mayıs|Haziran|Temmuz|"
            r"Ağustos|Eylül|Ekim|Kasım|Aralık))(\s*[-–]\s*)"
            r"(\d{1,2}\s+(?:Ocak|Şubat|Mart|Nisan|Mayıs|Haziran|Temmuz|"
            r"Ağustos|Eylül|Ekim|Kasım|Aralık))",
            lambda m: _gun_ay_metni(bas) + m.group(2) + _gun_ay_metni(bit),
            yeni)

        # 2) '01.01.2026-14.08.2026' bicimli donemler (ayni yil icinde)
        yeni = re.sub(
            r"01\.01\.%d\s*[-–]\s*\d{2}\.\d{2}\.%d" % (bas.year, bas.year),
            donem_nokta, yeni)

        # 3) Cari yil sayilari - sadece kesin taninan cumlelerde
        yeni = _sayilari_guncelle(yeni, d, degisen)

        if yeni != metin:
            if _paragraf_metni_degistir(p, yeni):
                degisen.append({
                    "paragraf": sira,
                    "onceki": metin.strip(),
                    "sonraki": yeni.strip(),
                    "farklar": _farklari_bul(metin, yeni),
                })

    # SAP'tan turetilemeyen sayilari kullaniciya bildir
    elle = _elle_kontrol_listesi(belge, bas)

    try:
        belge.save(docx_yolu)
    except Exception as e:
        raise WordHatasi(
            "Word dosyasi KAYDEDILEMEDI: %s\n"
            "Dosya Word'de acik olabilir; kapatip tekrar deneyin.\n"
            "Yedeginiz duruyor: %s" % (_kisa(e), yedek))

    if log:
        if degisen:
            log("  %d paragraf guncellendi:" % len(degisen))
            for x in degisen:
                ozet = x["sonraki"][:64].replace("\n", " ")
                log("    p%-4d %s..." % (x["paragraf"], ozet))
        else:
            log("  Degisiklik gerekmedi (belge zaten guncel).")
    return {"degisen": degisen, "elle": elle, "yedek": yedek,
            "excel_eksik": excel_eksik}


def _farklari_bul(onceki, sonraki):
    """
    Iki metin arasindaki degisen parcalari '30.06.2026 -> 17.08.2026'
    seklinde listeler. Kullanici cikitida NE degistigini gorsun diye.
    """
    import difflib
    farklar = []
    sm = difflib.SequenceMatcher(None, onceki, sonraki)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            continue
        # Degisen bolgeyi kelime sinirina genislet (anlamli gorunsun)
        a = onceki[max(0, i1 - 12):min(len(onceki), i2 + 12)].strip()
        b = sonraki[max(0, j1 - 12):min(len(sonraki), j2 + 12)].strip()
        if a != b:
            farklar.append((a, b))
    return farklar


def _sayilari_guncelle(metin, d, degisen):
    """
    Cumleyi tanidiktan SONRA sayiyi degistirir. Boylece ayni sayi
    baska bir baglamda geciyorsa yanlislikla degismez.
    """
    yeni = metin

    # a) "... döneminde ise 66.411 web sitesi tespit edilerek ..."
    if "web sitesi tespit edilerek" in metin or "web sitesi tespit edilerek" in yeni:
        yeni = re.sub(
            r"(döneminde ise\s+)([\d.]+)(\s+web\s+sitesi)",
            lambda m: m.group(1) + bicimle(d["toplam"]) + m.group(3),
            yeni)

    # b) Ülke dagilimi cumlesi: ilk 10 ulke yuzdesi + kalan
    if "hizmet aldıkları ülke dağılımı" in metin:
        yeni = _ulke_cumlesi_guncelle(yeni, d)

    # c) Calisma kitabindan gelen sayilar
    yeni = _excel_sayilari_guncelle(yeni, d.get("excel") or {})

    return yeni


# Cumleyi tanidiktan sonra icindeki sayiyi degistiren kurallar.
# Her kural: (excel_anahtari, desen, aciklama)
# Desende \\1 oncesi, \\2 sayi, \\3 sonrasi olacak sekilde gruplanir.
EXCEL_KURALLARI = [
    ("kumulatif_tespit",
     r"(döneminde toplam\s+)([\d.]+)(\s+web\s+sitesi\s+için\s+tespit)"),
    ("kumulatif_engel",
     r"(bunlardan\s+)([\d.]+)(['’]\w*\s+erişime\s+engellenmesi)"),
    ("mobil",
     r"(toplam\s+)([\d.]+)(\s+adet\s+mobil\s+uygulama)"),
    ("turetilmis_site",
     r"(kök\s+siteden\s+türetilmiş\s+)([\d.]+)(\s+internet\s+sitesi)"),
    ("kok_site",
     r"(toplam\s+)([\d.]+)(\s+kök\s+siteden)"),
    ("hat_850",
     r"(yılında\s+)([\d.]+)(\s+hattın)"),
    ("iban",
     r"(tespiti\s+sağlanan\s+)([\d.]+)(\s+banka/ödeme)"),
]


def _excel_sayilari_guncelle(metin, excel):
    """
    Calisma kitabindan okunan sayilari ilgili cumlelere yazar.

    Once cumle desenle taninir, SONRA icindeki sayi degistirilir; boylece
    ayni sayi baska bir baglamda geciyorsa yanlislikla degismez.
    Excel'den okunamayan deger icin cumleye DOKUNULMAZ.
    """
    if not excel:
        return metin
    yeni = metin
    for anahtar, desen in EXCEL_KURALLARI:
        deger = excel.get(anahtar)
        if deger is None:
            continue
        yeni = re.sub(desen,
                      lambda m: m.group(1) + bicimle(deger) + m.group(3),
                      yeni)
    return yeni


def _ulke_cumlesi_guncelle(metin, d):
    """
    '<Ulke>'nin %XX, <Ulke>'nin %XX, ...' bicimindeki cumledeki yuzdeleri
    SAP verisine gore gunceller.

    Ulke ADLARI degismez (cumle yapisi korunur); sadece ulke adindan
    hemen sonra gelen %XX degeri guncellenir. Boylece cumlenin dili
    ve sirasi bozulmaz.
    """
    yuzdeler, kalan = _ulke_yuzdeleri(d["ulkeler"], 10)
    if not yuzdeler:
        return metin

    yeni = metin
    for ad, yuzde in yuzdeler:
        # Ulke adi -> belgede '<Ulke>’nin %XX' bicimindedir
        kok = re.escape(ad.split("(")[0].strip())
        # Ulke adi + ekler + %sayi
        yeni = re.sub(
            r"(%s[’'']?\w*\s*%%)\s*\d+(?:,\d+)?" % kok,
            lambda m: m.group(1) + str(yuzde),
            yeni, count=1, flags=re.IGNORECASE)

    # 'kalan %8'lik kısmın' / 'kalan %3lük dilimin'
    yeni = re.sub(r"(kalan\s*%)\s*\d+(?:,\d+)?", r"\g<1>" + str(kalan), yeni)
    return yeni


# Programin GUNCELLEMEDIGI sayilar: hangi cumle, nerede duruyor.
#   (word_anahtari, aciklama, nerede_duzeltilir)
ELLE_KALEMLER = [
    ("erişime engellenmesi sağlanmıştır",
     "'erişime engellenen' kümülatif sayısı",
     "İhbar Site Rapor > Yıllar Bazlı > E23-I23 (cari yıl satırı)"),
    ("Reklam Kuruluna bildirilerek",
     "Reklam Kurulu'na bildirilen hesap sayıları",
     "kaynağı yok - Word'de elle yazılır"),
    ("internet sitesi için yetkili Cumhuriyet",
     "suç duyurusu yıl bazlı site sayıları (2024/2025/2026)",
     "İhbar Site Rapor > Suç Duyuruları > yeni dönem satırı eklenmeli"),
    ("sosyal medya hesabı için yetkili",
     "2025 yılı sosyal medya sayıları (2026 otomatik)",
     "kaynağı yok - Word'de elle yazılır"),
]


def _elle_kontrol_listesi(belge, bas):
    """
    Programin GUNCELLEMEDIGI sayilari, NEREDE duzeltilecegiyle birlikte
    listeler. Boylece kullanici hangi dosyanin hangi sekmesine gidecegini
    bilir.

    Doner: [(aciklama, nerede), ...]
    """
    bulunan = []
    gorulen = set()
    for p in belge.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        for anahtar, aciklama, nerede in ELLE_KALEMLER:
            if anahtar in t and aciklama not in gorulen:
                gorulen.add(aciklama)
                bulunan.append((aciklama, nerede))
                break
    return bulunan
